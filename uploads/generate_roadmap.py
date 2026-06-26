from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── Colour palette ────────────────────────────────────────────────────────────
DARK_NAVY  = RGBColor(0x0D, 0x1B, 0x2A)
BRAND_BLUE = RGBColor(0x00, 0x6E, 0xFF)
ACCENT     = RGBColor(0x00, 0xC2, 0xA8)
LIGHT_BG   = RGBColor(0xF0, 0xF4, 0xFF)
WHITE      = RGBColor(0xFF, 0xFF, 0xFF)
BODY_GRAY  = RGBColor(0x33, 0x33, 0x44)
MID_GRAY   = RGBColor(0x88, 0x88, 0x99)

# ── Helper: set paragraph shading ─────────────────────────────────────────────
def shade_paragraph(para, hex_fill):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_fill)
    pPr.append(shd)

# ── Helper: set table cell background ─────────────────────────────────────────
def cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

# ── Helper: set table cell borders ────────────────────────────────────────────
def cell_borders(cell, color_hex='D0D8F0', sz='4'):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ('top','left','bottom','right'):
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'),   'single')
        b.set(qn('w:sz'),    sz)
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), color_hex)
        tcBorders.append(b)
    tcPr.append(tcBorders)

# ── Helper: add a run with full formatting ────────────────────────────────────
def styled_run(para, text, bold=False, italic=False, size=11,
               color=None, font='Calibri'):
    run = para.add_run(text)
    run.bold   = bold
    run.italic = italic
    run.font.name = font
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return run

# ══════════════════════════════════════════════════════════════════════════════
#  COVER BANNER
# ══════════════════════════════════════════════════════════════════════════════
banner = doc.add_paragraph()
banner.alignment = WD_ALIGN_PARAGRAPH.CENTER
shade_paragraph(banner, '006EFF')
banner.paragraph_format.space_before = Pt(18)
banner.paragraph_format.space_after  = Pt(4)

styled_run(banner, 'ALTMATIC ADMS', bold=True, size=28,
           color=WHITE, font='Calibri')

sub_banner = doc.add_paragraph()
sub_banner.alignment = WD_ALIGN_PARAGRAPH.CENTER
shade_paragraph(sub_banner, '006EFF')
sub_banner.paragraph_format.space_before = Pt(0)
sub_banner.paragraph_format.space_after  = Pt(18)
styled_run(sub_banner, '30-Day Launch Roadmap', bold=False, size=14,
           color=RGBColor(0xCC, 0xE0, 0xFF), font='Calibri')

# tagline
tagline = doc.add_paragraph()
tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
shade_paragraph(tagline, 'F0F4FF')
tagline.paragraph_format.space_before = Pt(0)
tagline.paragraph_format.space_after  = Pt(16)
styled_run(tagline, 'Build  ·  Market  ·  Sell  —  Your Automated Digital Marketing System',
           italic=True, size=11, color=BRAND_BLUE, font='Calibri')

# ══════════════════════════════════════════════════════════════════════════════
#  SECTION HEADER HELPER
# ══════════════════════════════════════════════════════════════════════════════
def section_header(title, subtitle='', phase_color='006EFF'):
    hdr = doc.add_paragraph()
    hdr.alignment = WD_ALIGN_PARAGRAPH.LEFT
    shade_paragraph(hdr, phase_color)
    hdr.paragraph_format.space_before = Pt(14)
    hdr.paragraph_format.space_after  = Pt(2)
    hdr.paragraph_format.left_indent  = Cm(0.3)
    styled_run(hdr, f'  {title}', bold=True, size=15, color=WHITE, font='Calibri')
    if subtitle:
        sub = doc.add_paragraph()
        shade_paragraph(sub, phase_color)
        sub.alignment = WD_ALIGN_PARAGRAPH.LEFT
        sub.paragraph_format.space_before = Pt(0)
        sub.paragraph_format.space_after  = Pt(10)
        sub.paragraph_format.left_indent  = Cm(0.3)
        styled_run(sub, f'  {subtitle}', bold=False, italic=True, size=10,
                   color=WHITE, font='Calibri')

def week_header(title):
    wk = doc.add_paragraph()
    wk.alignment = WD_ALIGN_PARAGRAPH.LEFT
    shade_paragraph(wk, 'E8F0FF')
    wk.paragraph_format.space_before = Pt(10)
    wk.paragraph_format.space_after  = Pt(4)
    wk.paragraph_format.left_indent  = Cm(0.3)
    styled_run(wk, f'  {title}', bold=True, size=11,
               color=BRAND_BLUE, font='Calibri')

# ══════════════════════════════════════════════════════════════════════════════
#  TABLE BUILDER
# ══════════════════════════════════════════════════════════════════════════════
def add_table(headers, rows, col_widths, header_color='0D1B2A', stripe_color='F7F9FF'):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = 'Table Grid'

    # header row
    hdr_row = tbl.rows[0]
    for i, (cell, txt) in enumerate(zip(hdr_row.cells, headers)):
        cell_bg(cell, header_color)
        cell_borders(cell, 'FFFFFF', '6')
        cell.width = Inches(col_widths[i])
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after  = Pt(4)
        styled_run(p, txt, bold=True, size=10, color=WHITE, font='Calibri')

    # data rows
    for r_idx, row_data in enumerate(rows):
        row = tbl.rows[r_idx + 1]
        bg  = stripe_color if r_idx % 2 == 0 else 'FFFFFF'
        for c_idx, (cell, txt) in enumerate(zip(row.cells, row_data)):
            cell_bg(cell, bg)
            cell_borders(cell, 'D0D8F0', '4')
            cell.width = Inches(col_widths[c_idx])
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after  = Pt(3)
            align = WD_ALIGN_PARAGRAPH.CENTER if c_idx == 0 else WD_ALIGN_PARAGRAPH.LEFT
            p.alignment = align
            is_bold  = (c_idx == 0)
            txt_color = BRAND_BLUE if c_idx == 0 else BODY_GRAY
            styled_run(p, txt, bold=is_bold, size=10,
                       color=txt_color, font='Calibri')

    doc.add_paragraph().paragraph_format.space_after = Pt(4)

# ══════════════════════════════════════════════════════════════════════════════
#  PHASE 1 — BUILD
# ══════════════════════════════════════════════════════════════════════════════
section_header('PHASE 1 — BUILD', 'Days 1–10  |  Core System + Productize', '0D1B2A')

week_header('Week 1 — Core System')
add_table(
    headers=['Day', 'Task', 'Output'],
    col_widths=[0.55, 2.8, 3.15],
    rows=[
        ['1',  'Define your niche & ICP (Ideal Customer Profile)',
               '1-page brief: who it\'s for, what pain it solves'],
        ['2',  'Map the automation workflows',
               'Flowchart: lead capture → nurture → convert → retain'],
        ['3',  'Set up tech stack (CRM, email, landing page, scheduler)',
               'Live integrations (Make.com / n8n + Mailchimp + Webflow)'],
        ['4',  'Build lead capture + welcome sequence',
               'Opt-in page + 3-email welcome flow'],
        ['5',  'Build nurture sequence',
               '5–7 email series, segmented by interest'],
        ['6',  'Build conversion sequence',
               'Sales email series + offer page'],
        ['7',  'QA all workflows end-to-end',
               'Zero broken triggers, correct data passing'],
    ]
)

week_header('Week 2 — Productize It')
add_table(
    headers=['Day', 'Task', 'Output'],
    col_widths=[0.55, 2.8, 3.15],
    rows=[
        ['8',  'Package the system (tiers/pricing)',
               'Starter / Growth / Done-For-You pricing page'],
        ['9',  'Build onboarding + delivery process',
               'Client intake form + SOP document'],
        ['10', 'Record a 3-minute demo video',
               'Loom walkthrough of the live system'],
    ]
)

# ══════════════════════════════════════════════════════════════════════════════
#  PHASE 2 — MARKET
# ══════════════════════════════════════════════════════════════════════════════
section_header('PHASE 2 — MARKET', 'Days 11–20  |  Audience, Content & Outreach', '006EFF')

week_header('Weeks 2–3 — Audience & Content')
add_table(
    headers=['Day', 'Task', 'Output'],
    col_widths=[0.55, 2.8, 3.15],
    rows=[
        ['11', 'Write your core positioning statement',
               '1 sentence: who, what, result, timeframe'],
        ['12', 'Set up LinkedIn + one social channel',
               'Optimised profile + pinned content'],
        ['13', 'Publish your origin story post',
               '"Why I built this" — humanises the offer'],
        ['14', 'Launch a free resource (lead magnet)',
               'PDF checklist / mini-course / audit template'],
        ['15', 'Set up your own ADMS for marketing',
               'Eat your own cooking — automate your outreach'],
        ['16', 'Publish 3 problem-aware posts',
               'Address top 3 pain points of your ICP'],
        ['17', 'Start cold outreach — 10 ideal prospects/day',
               'Personalised DMs or emails'],
        ['18', 'Host a free 30-min live demo or webinar',
               'Zoom/Loom — show the system live'],
        ['19', 'Collect 3 testimonials / case studies',
               'Beta user quotes or your own results'],
        ['20', 'Launch a referral loop',
               '"Refer 1 client, get 1 month free"'],
    ]
)

# ══════════════════════════════════════════════════════════════════════════════
#  PHASE 3 — SELL
# ══════════════════════════════════════════════════════════════════════════════
section_header('PHASE 3 — SELL', 'Days 21–30  |  Pipeline, Proposals & Close', '00C2A8')

week_header('Week 4 — Pipeline & Close')
add_table(
    headers=['Day', 'Task', 'Output'],
    col_widths=[0.55, 2.8, 3.15],
    rows=[
        ['21', 'Set up your sales pipeline in CRM',
               'Stages: Lead → Demo → Proposal → Close'],
        ['22', 'Write sales script / discovery call framework',
               '5-question discovery + objection handlers'],
        ['23', 'Run first paid discovery calls',
               '3–5 calls booked'],
        ['24', 'Send proposals to qualified leads',
               'Templated proposal + e-sign (DocuSign / PandaDoc)'],
        ['25', 'Follow-up sequence activated',
               '3-touch follow-up automation'],
        ['26', 'Close first 3 paying clients',
               'Collect payment + start onboarding'],
        ['27', 'Gather feedback from clients',
               '1:1 call or async Loom'],
        ['28', 'Document wins + iterate offer',
               'Update pricing / positioning if needed'],
        ['29', 'Launch a public case study',
               'Post results with permission'],
        ['30', 'Plan Month 2 — scale what worked',
               'Double down on best-performing channel'],
    ]
)

# ══════════════════════════════════════════════════════════════════════════════
#  KEY METRICS
# ══════════════════════════════════════════════════════════════════════════════
section_header('KEY METRICS TO TRACK', 'Target by Day 30', '0D1B2A')
add_table(
    headers=['Metric', 'Target by Day 30'],
    col_widths=[3.3, 3.2],
    header_color='0D1B2A',
    rows=[
        ['Leads captured',    '100+'],
        ['Demo calls booked', '10+'],
        ['Proposals sent',    '5+'],
        ['Clients closed',    '3+'],
        ['MRR generated',     '$1,500 – $5,000+'],
    ]
)

# ══════════════════════════════════════════════════════════════════════════════
#  TECH STACK
# ══════════════════════════════════════════════════════════════════════════════
section_header('RECOMMENDED TECH STACK', '', '006EFF')
add_table(
    headers=['Category', 'Tool Options'],
    col_widths=[2.0, 4.5],
    header_color='006EFF',
    rows=[
        ['Automation',    'Make.com  |  n8n'],
        ['CRM',           'GoHighLevel  |  HubSpot Free  |  Notion CRM'],
        ['Email',         'Mailchimp  |  ActiveCampaign  |  ConvertKit'],
        ['Landing Pages', 'Webflow  |  Carrd  |  GoHighLevel'],
        ['Scheduling',    'Calendly'],
        ['Payments',      'Stripe'],
        ['Proposals',     'PandaDoc  |  Notion'],
    ]
)

# ══════════════════════════════════════════════════════════════════════════════
#  FOOTER
# ══════════════════════════════════════════════════════════════════════════════
footer_p = doc.add_paragraph()
footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
shade_paragraph(footer_p, '0D1B2A')
footer_p.paragraph_format.space_before = Pt(20)
footer_p.paragraph_format.space_after  = Pt(8)
styled_run(footer_p, 'Altmatic ADMS  ·  Automated Digital Marketing System  ·  2026',
           size=9, color=MID_GRAY, font='Calibri')

# ── Save ───────────────────────────────────────────────────────────────────────
output_path = '/Users/jaysingh/Desktop/AltmaticADMS/Altmatic_ADMS_30Day_Roadmap.docx'
doc.save(output_path)
print(f'Saved → {output_path}')
